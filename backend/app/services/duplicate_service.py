import os
import tempfile
from pathlib import Path
from typing import List, Dict, Set
from app.utils.logger import logger
import docx as dx
import re
import time

class DuplicateService:
    def __init__(self):
        self.task_results: Dict[str, dict] = {}
        self.temp_dirs: Dict[str, str] = {}
    
    def _save_files(self, files: List, task_id: str) -> List[str]:
        """保存上传文件到临时目录"""
        # 使用系统临时目录
        temp_base = Path(tempfile.gettempdir()) / "bid_check"
        temp_base.mkdir(exist_ok=True)
        
        task_dir = temp_base / task_id
        task_dir.mkdir(exist_ok=True)
        
        paths = []
        for file in files:
            # 验证文件格式
            if not file.filename.lower().endswith('.docx'):
                raise ValueError(f"不支持的文件格式: {file.filename}")
            
            file_path = task_dir / file.filename
            
            try:
                # 检查文件是否已关闭
                if hasattr(file.file, 'closed') and file.file.closed:
                    # 如果文件已关闭，需要重新获取文件内容，但这在FastAPI中不可行
                    # 所以我们尝试先读取内容再处理
                    raise ValueError(f"文件句柄已关闭: {file.filename}")
                
                # 读取文件内容到内存
                file.file.seek(0)  # 重置文件指针
                content = file.file.read()
                
                # 限制文件大小
                if len(content) > 50 * 1024 * 1024:  # 50MB限制
                    raise ValueError(f"文件过大: {file.filename}")
                
                # 写入文件
                with open(file_path, "wb") as f:
                    f.write(content)
                
                # 验证文件是否成功保存
                if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
                    raise ValueError(f"文件保存失败: {file.filename}")
                
                paths.append(str(file_path))
                logger.info(f"文件保存成功: {file.filename} -> {file_path}")
                
            except ValueError:
                # 重新抛出值错误，不进行额外处理
                raise
            except Exception as e:
                logger.error(f"保存文件失败 {file.filename}: {str(e)}")
                # 清理已保存的文件
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise
        
        return paths
    
    def _get_sentence(self, doc, splitword="", limitnum=10, needrun=False):
        """获取文档中的所有短句"""
        for p in doc.paragraphs:
            for run in p.runs:
                text = run.text
                text = text.replace(" ", "").replace("\t", "").replace("\r", "").replace("\n", "")
                if len(splitword) > 0:
                    for t in re.split(splitword, text):
                        if len(t) >= limitnum:
                            if needrun:
                                yield [t, run]
                            else:
                                yield t
                else:
                    if needrun:
                        yield [text, run]
                    else:
                        yield text
        for table in doc.tables:
            for row in table.rows:
                try:
                    row.cells[0].text
                except:
                    continue
                for cell in row.cells:
                    text = cell.text
                    text = text.replace(" ", "").replace("\t", "").replace("\r", "").replace("\n", "")
                    if len(splitword) > 0:
                        for t in re.split(splitword, text):
                            if len(t) >= limitnum:
                                if needrun:
                                    class VirtualRun:
                                        def __init__(self, text):
                                            self.text = text
                                            self.font = type('obj', (object,), {
                                                'color': type('obj', (object,), {'rgb': None})
                                            })()
                                    yield [t, VirtualRun(t)]
                                else:
                                    yield t
                    else:
                        if needrun:
                            class VirtualRun:
                                def __init__(self, text):
                                    self.text = text
                                    self.font = type('obj', (object,), {
                                        'color': type('obj', (object,), {'rgb': None})
                                    })()
                            yield [text, VirtualRun(text)]
                        else:
                            yield text
    
    def _compare(self, files, limitnum, splitword):
        """核心对比算法 - 只对比投标文件之间的重复"""
        try:
            starttime = time.time()
            
            # 获取投标文档内容
            logger.info(f"开始读取投标文件，共{len(files)}个")
            text_list = []
            doc_all = []
            for i, file in enumerate(files):
                logger.info(f"处理投标文件 {i+1}/{len(files)}: {os.path.basename(file)}")
                try:
                    if not os.path.exists(file):
                        raise FileNotFoundError(f"文件不存在: {file}")
                    
                    text_set = set()
                    with open(file, 'rb') as f:
                        doc = dx.Document(f)
                        doc_all.append(doc)
                        for t in self._get_sentence(doc, splitword, limitnum):
                            text_set.add(t)
                    text_list.append(text_set)
                except Exception as e:
                    logger.error(f"读取投标文件失败 {file}: {str(e)}")
                    raise
            
            readtime = time.time()
            logger.info(f"文件读取完成,用时{readtime - starttime:.2f}秒")
            
            # 记录详细统计信息
            logger.info(f"投标文件数量: {len(text_list)}")
            for i, text_set in enumerate(text_list):
                logger.info(f"投标文件 {i+1} 文本数量: {len(text_set)}")
            
            # 寻找重复字句（投标文件之间的重复）
            logger.info("开始查找重复字句（投标文件之间）")
            text_error_list = []
            for n in range(len(text_list)):
                text_error = set()
                text_n = text_list[n]
                total_checks = len(text_n)
                check_count = 0
                match_count = 0
                
                logger.info(f"文件 {n+1} 开始检查，共 {total_checks} 条文本")
                
                for text in text_n:
                    check_count += 1
                    if check_count % 1000 == 0 or check_count == total_checks:
                        logger.info(f"文件 {n+1} 检查进度: {check_count}/{total_checks}")
                    
                    if len(text) < limitnum:
                        continue
                    for i in range(len(text_list)):
                        if i != n:
                            if text in text_list[i]:
                                text_error.add(text)
                                match_count += 1
                                break
                
                logger.info(f"文件 {n+1} 检查完成，发现 {len(text_error)} 条重复")
                text_error_list.append(text_error)
            
            errortime = time.time()
            logger.info(f"寻找重复字句完毕，用时{errortime - readtime:.2f}秒")
            
            # 生成结果
            results = []
            for i in range(len(files)):
                file_name = os.path.basename(files[i])
                duplicate_count = len(text_error_list[i])
                total_count = len(text_list[i])
                duplicate_rate = (duplicate_count / total_count) * 100 if total_count > 0 else 0
                
                logger.info(f"文件 {i+1} 重复率详细计算:")
                logger.info(f"  - 总文本数量 (total_count): {total_count}")
                logger.info(f"  - 重复文本数量 (duplicate_count): {duplicate_count}")
                logger.info(f"  - 重复率: {duplicate_rate:.2f}%")
                
                # 确保重复率不会超过100%（防止计算错误导致的异常值）
                duplicate_rate = min(duplicate_rate, 100.0)
                duplicate_count = min(duplicate_count, total_count)
                
                results.append({
                    "file_name": file_name,
                    "duplicate_rate": round(duplicate_rate, 2),
                    "duplicate_count": duplicate_count,
                    "total_count": total_count,
                    "duplicate_sections": list(text_error_list[i])
                })
            
            return results
        except Exception as e:
            logger.error(f"查重算法执行失败: {str(e)}", exc_info=True)
            raise
    
    async def check_duplicate(
        self, 
        task_id: str, 
        tender_files: List, 
        bid_files: List, 
        params
    ):
        """执行查重任务"""
        try:
            # 设置初始状态
            self.task_results[task_id] = {
                "status": "processing",
                "results": None
            }
            
            # 保存文件
            tender_paths = self._save_files(tender_files, task_id)
            bid_paths = self._save_files(bid_files, task_id)
            
            # 执行查重
            results = self._compare(
                bid_paths, 
                params.min_length, 
                params.split_words
            )
            
            # 保存结果
            self.task_results[task_id] = {
                "status": "completed",
                "results": results
            }
            
            # 清理临时文件
            for path in tender_paths + bid_paths:
                os.remove(path)
                
        except Exception as e:
            logger.error(f"查重失败: {str(e)}")
            self.task_results[task_id] = {
                "status": "failed",
                "error": str(e)
            }
    
    async def check_duplicate_with_content(
        self, 
        task_id: str, 
        bid_file_contents: List[dict], 
        params
    ):
        """使用预读取的文件内容执行查重任务（解决文件句柄关闭问题）"""
        import io
        import shutil
        logger.info(f"[{task_id}] 开始执行check_duplicate_with_content方法")
        logger.info(f"[{task_id}] 投标文件数量: {len(bid_file_contents)}")
        
        try:
            # 设置初始状态
            self.task_results[task_id] = {
                "status": "processing",
                "results": None
            }
            
            # 使用系统临时目录
            temp_base = Path(tempfile.gettempdir()) / "bid_check"
            temp_base.mkdir(exist_ok=True)
            
            task_dir = temp_base / task_id
            task_dir.mkdir(exist_ok=True)
            
            bid_paths = []
            for i, file_info in enumerate(bid_file_contents):
                filename = file_info["filename"]
                content = file_info["content"]
                logger.info(f"[{task_id}] 处理投标文件 {i+1}: {filename}, 内容大小: {len(content)} bytes")
                file_path = task_dir / filename
                
                # 验证文件格式
                if not filename.lower().endswith('.docx'):
                    raise ValueError(f"不支持的文件格式: {filename}")
                
                # 验证文件大小
                if len(content) > 50 * 1024 * 1024:  # 50MB限制
                    raise ValueError(f"文件过大: {filename}")
                
                # 写入文件
                with open(file_path, "wb") as f:
                    f.write(content)
                
                # 验证文件是否成功保存
                if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
                    raise ValueError(f"文件保存失败: {filename}")
                
                bid_paths.append(str(file_path))
                logger.info(f"[{task_id}] 投标文件保存成功: {filename} -> {file_path}")
            
            logger.info(f"[{task_id}] 开始执行查重算法")
            
            # 执行查重
            results = self._compare(
                bid_paths, 
                params.min_length, 
                params.split_words
            )
            
            logger.info(f"[{task_id}] 查重完成，结果数量: {len(results)}")
            
            # 保存结果
            self.task_results[task_id] = {
                "status": "completed",
                "results": results
            }
            
            # 清理临时文件
            for path in bid_paths:
                if os.path.exists(path):
                    os.remove(path)
            
            # 清理任务目录
            if os.path.exists(task_dir):
                shutil.rmtree(task_dir)
                
        except Exception as e:
            logger.error(f"[{task_id}] 查重失败: {str(e)}", exc_info=True)
            self.task_results[task_id] = {
                "status": "failed",
                "error": str(e)
            }
    
    def get_result(self, task_id: str) -> dict:
        """获取任务结果"""
        result = self.task_results.get(task_id, {})
        if not result:
            return {"task_id": task_id, "status": "not_found"}
        
        # 确保返回的结果包含task_id字段
        result_with_id = result.copy()
        result_with_id["task_id"] = task_id
        return result_with_id